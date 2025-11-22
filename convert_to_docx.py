"""
Convert instructions.md to instructions.docx
Simple markdown to Word document converter
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def parse_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to DOCX"""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Title (first line)
        if i == 0:
            p = doc.add_heading(line, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Subtitle (second line)
        if i == 1:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(14)
            p.runs[0].font.color.rgb = RGBColor(0, 120, 212)
            i += 1
            continue

        # Main headings
        if not line.startswith(' ') and len(line) > 0 and i > 1:
            # Check if this is a heading by looking at the next line
            if i + 1 < len(lines) and not lines[i + 1].strip():
                doc.add_heading(line, level=1)
                i += 1
                continue

        # Subheadings (sections that appear to be headings)
        if line and not line.startswith(' ') and not line.startswith('-') and not line.startswith('*'):
            # Check if followed by content (making it likely a heading)
            if i + 1 < len(lines) and lines[i + 1].strip() and not lines[i + 1].startswith(' '):
                doc.add_heading(line, level=2)
                i += 1
                continue

        # Code blocks
        if line.strip().startswith('{') or 'pip install' in line or 'python' in line or line.strip().startswith('http'):
            p = doc.add_paragraph(line)
            p.style = 'Intense Quote'
            run = p.runs[0]
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            i += 1
            continue

        # Bullet points or lists
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        # Regular paragraphs
        if line.strip():
            p = doc.add_paragraph(line)
            i += 1
        else:
            i += 1

    # Save document
    doc.save(docx_file)
    print(f"Converted {md_file} to {docx_file}")

if __name__ == "__main__":
    parse_markdown_to_docx('instructions.md', 'instructions.docx')
