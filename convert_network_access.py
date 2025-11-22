"""
Convert NETWORK_ACCESS.md to clean format without markdown symbols
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def clean_markdown_line(line):
    """Remove markdown formatting from a line"""
    # Remove heading markers
    line = re.sub(r'^#{1,6}\s+', '', line)
    # Remove bold/italic markers
    line = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)
    line = re.sub(r'\*([^*]+)\*', r'\1', line)
    # Remove inline code markers
    line = re.sub(r'`([^`]+)`', r'\1', line)
    return line

def convert_network_access():
    """Convert NETWORK_ACCESS.md to docx with clean formatting"""

    # Read the markdown file
    with open('NETWORK_ACCESS.md', 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Create Word document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    i = 0
    in_code_block = False

    while i < len(lines):
        line = lines[i].rstrip()
        original_line = line

        # Handle code blocks
        if line.startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            # Add code as monospace
            if line.strip():
                p = doc.add_paragraph(line)
                run = p.runs[0] if p.runs else p.add_run(line)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            i += 1
            continue

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Title (# Network Access Guide)
        if line.startswith('# '):
            title_text = clean_markdown_line(line)
            p = doc.add_heading(title_text, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Main sections (## )
        if line.startswith('## '):
            heading_text = clean_markdown_line(line)
            doc.add_heading(heading_text, level=1)
            i += 1
            continue

        # Subsections (### )
        if line.startswith('### '):
            heading_text = clean_markdown_line(line)
            doc.add_heading(heading_text, level=2)
            i += 1
            continue

        # Tables - detect and skip markdown tables
        if '|' in line and (line.count('|') >= 3 or (i+1 < len(lines) and '---' in lines[i+1])):
            # Skip table separator line
            if '---' in line:
                i += 1
                continue
            # Format table rows as simple text
            parts = [p.strip() for p in line.split('|') if p.strip()]
            if parts:
                doc.add_paragraph(' | '.join(parts))
            i += 1
            continue

        # Bullet points
        if line.startswith('- ') or line.startswith('* '):
            text = clean_markdown_line(line[2:].strip())
            doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        # Regular paragraphs
        if line.strip():
            clean_line = clean_markdown_line(line)
            doc.add_paragraph(clean_line)

        i += 1

    # Save document
    doc.save('NETWORK_ACCESS.docx')
    print("Created NETWORK_ACCESS.docx")

if __name__ == "__main__":
    convert_network_access()
